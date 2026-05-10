# build_submission_docx.py
# Auto-assemble HOMEWORK3_submission.docx with rubric table, embedded stats, sample report,
# CSV head, and two figures. The grader sees a near-complete document; the student only needs
# to write Section 1 (NOT AI-generated, ~500 words) and add their own GitHub link if forked.
#
# Usage (after run_experiment.py + analyze_experiment.py have produced output/):
#   python build_submission_docx.py

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


PROMPT_ORDER = ["A_minimal_paragraph", "B_structured_sections", "C_ICS_traceability"]


def _add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_df_table(doc: Document, df: pd.DataFrame, max_rows: int = 12) -> None:
    df = df.head(max_rows)
    cols = list(df.columns)
    table = doc.add_table(rows=len(df) + 1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = c
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        for c_idx, c in enumerate(cols):
            v = row[c]
            text = "" if pd.isna(v) else str(v)
            table.rows[r_idx].cells[c_idx].text = text


def main() -> None:
    root = Path(__file__).resolve().parent
    out_dir = root / "output"
    docx_path = root / "HOMEWORK3_submission.docx"
    csv_path = out_dir / "validation_results.csv"
    stats_path = out_dir / "stats_summary.txt"
    perdim_path = out_dir / "per_dimension_summary.csv"
    box_png = out_dir / "boxplot_composite.png"
    bar_png = out_dir / "bar_per_dimension.png"
    meta_path = out_dir / "run_metadata.json"

    doc = Document()
    h = doc.add_heading("Homework 3 — AI Report Validation System", level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "Student name: _____________________      Date: _____________________"
    )

    # --- Section 1: Writing (NOT AI-generated) ---
    doc.add_heading("1. Written explanation (~500 words) — MUST be written by the student", level=1)
    instr = doc.add_paragraph()
    instr.add_run(
        "Replace this paragraph with ~500 words in your own words. Cover: "
        "(a) purpose of the validation system; (b) how your custom rubric differs from the LAB Likert scales; "
        "(c) experimental design — three prompts (A/B/C), replicates per prompt, total N; "
        "(d) statistical results — copy F and p_anova from Section 4, name the winning prompt and Cohen's d; "
        "(e) design choices and challenges. Do NOT use an AI to write this section."
    ).italic = True

    # --- Section 2: Git links ---
    doc.add_heading("2. Git repository links", level=1)
    import os as _os
    base = _os.environ.get(
        "HW3_REPO_BASE_URL",
        "https://github.com/zj276-commits/Homework/blob/main",
    )
    for label, tail in [
        ("Validation runner (generation + AI reviewer)", "run_experiment.py"),
        ("Statistical analysis (ANOVA + t-tests + plots)", "analyze_experiment.py"),
        ("Custom rubric source / scenario fact pack", "data/scenario_source.md"),
        ("Example validation outputs", "output/validation_results.csv"),
        ("Per-dimension summary table", "output/per_dimension_summary.csv"),
        ("Full documentation", "README.md"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(f"{base}/{tail}")
    doc.add_paragraph(
        "[Optional] If you reused Homework 1/2 reports as validator inputs (via --reports-dir), "
        "add links to those report files here as well."
    )

    # --- Section 3: Validation rubric table ---
    doc.add_heading("3. Validation rubric (customized — not LAB Likert scales)", level=1)
    rubric = doc.add_table(rows=6, cols=4)
    rubric.style = "Light Grid Accent 1"
    hdr = ["Dimension", "Description", "Scale", "Contrast with LAB Likert (02_ai_quality_control.R)"]
    for i, t in enumerate(hdr):
        rubric.rows[0].cells[i].text = t
    rubric_rows = [
        ("scenario_fidelity_pct", "Coverage of labeled facts F1–F10 without contradiction.",
         "0–100 %", "LAB scores generic 'accuracy' on a 1–5 Likert; we tie it to a labeled fact pack."),
        ("decision_actionability_pct", "Can a coordinator pick next-hour actions from the report alone?",
         "0–100 %", "LAB has no 'usability' axis; this is operational, not literary."),
        ("briefing_structure_score", "EOC-scannable layout (sections, no wall of text).",
         "0–10 integer", "LAB measures 'clarity/formality' Likert; we score scannable structure directly."),
        ("traceability_pct", "Explicit linkage of claims to fact IDs (F1, F2, …) or equivalent.",
         "0–100 %", "LAB has no traceability dimension."),
        ("overreach_severity", "False real-world escalation, fabricated URLs, contradicting drill scope.",
         "0–4 integer (penalty)", "LAB uses one boolean 'accurate'; we encode degrees of drill breach."),
    ]
    for r, row in enumerate(rubric_rows, start=1):
        for c, val in enumerate(row):
            rubric.rows[r].cells[c].text = val
    doc.add_paragraph(
        "Composite score: equal-weighted mean of the 4 normalized positive dimensions "
        "(briefing_structure_score is multiplied by 10 to share a 0–100 scale). "
        "overreach_severity is analyzed separately."
    )

    # --- Section 4: Stats results (auto-embedded) ---
    doc.add_heading("4. Statistical analysis (auto-embedded from output/)", level=1)
    if stats_path.is_file():
        _add_code(doc, stats_path.read_text(encoding="utf-8"))
    else:
        doc.add_paragraph(
            "[stats_summary.txt not found — run: python analyze_experiment.py]"
        )

    if perdim_path.is_file():
        doc.add_heading("4a. Per-dimension means and ANOVA", level=2)
        _add_df_table(doc, pd.read_csv(perdim_path))
    else:
        doc.add_paragraph("[per_dimension_summary.csv not found — run analyze_experiment.py]")

    # --- Section 5: Figures ---
    doc.add_heading("5. Figures (auto-embedded)", level=1)
    if box_png.is_file():
        doc.add_paragraph("Figure 1. Composite score by prompt (boxplot).").bold = True
        doc.add_picture(str(box_png), width=Inches(5.8))
    else:
        doc.add_paragraph("[boxplot_composite.png missing — re-run analyze_experiment.py]")
    if bar_png.is_file():
        doc.add_paragraph("Figure 2. Per-dimension comparison across prompts.").bold = True
        doc.add_picture(str(bar_png), width=Inches(6.3))
    else:
        doc.add_paragraph("[bar_per_dimension.png missing — re-run analyze_experiment.py]")

    # --- Section 6: Sample report + sample validation row ---
    doc.add_heading("6. Sample inputs/outputs", level=1)
    sample_report = None
    reports_dir = out_dir / "reports"
    if reports_dir.is_dir():
        cands = sorted(reports_dir.glob("*.txt"))
        if cands:
            sample_report = cands[0]
    if sample_report:
        doc.add_paragraph(f"Sample generated report (`{sample_report.name}`):").bold = True
        _add_code(doc, sample_report.read_text(encoding="utf-8")[:2000])
    if csv_path.is_file():
        df = pd.read_csv(csv_path)
        keep_cols = [
            c for c in [
                "prompt_key", "replicate",
                "scenario_fidelity_pct", "decision_actionability_pct",
                "briefing_structure_score", "traceability_pct",
                "overreach_severity", "composite_score", "validator_notes",
            ] if c in df.columns
        ]
        doc.add_heading("Validation results — first rows of validation_results.csv", level=2)
        _add_df_table(doc, df[keep_cols].head(10))

    # --- Section 7: Run metadata + usage ---
    doc.add_heading("7. Reproducibility & usage", level=1)
    if meta_path.is_file():
        meta = json.loads(meta_path.read_text(encoding="utf-8"))
        doc.add_paragraph("Run metadata:").bold = True
        for k, v in meta.items():
            doc.add_paragraph(f"  {k}: {v}", style="List Bullet")
    _add_code(
        doc,
        "cd 11_decision_support/homework3\n"
        "py -3.12 -m venv .venv\n"
        ".venv\\Scripts\\pip install -r requirements.txt\n"
        "copy .env.example .env  (then set OLLAMA_API_KEY)\n"
        ".venv\\Scripts\\python run_experiment.py\n"
        ".venv\\Scripts\\python analyze_experiment.py\n"
        ".venv\\Scripts\\python build_submission_docx.py",
    )

    # --- Section 8: Screenshot checklist (for Canvas) ---
    doc.add_heading("8. Screenshot checklist (paste yours below)", level=1)
    for item in [
        "Terminal showing run_experiment.py producing 18+ rows (or --mock)",
        "One generated report .txt open in editor",
        "validation_results.csv open in Excel / VS Code",
        "stats_summary.txt or terminal output of analyze_experiment.py",
        "boxplot_composite.png (already embedded above; you can also screenshot it)",
        "bar_per_dimension.png (already embedded above)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")
    print("Open it in Word, write Section 1 yourself, paste screenshots, and submit on Canvas.")


if __name__ == "__main__":
    main()
